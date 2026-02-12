from docx import Document
from datetime import datetime, timedelta
from pathlib import Path

import os

class DocumentGenerator:
    def __init__(self, templates_dir="documents/templates", output_dir="documents/generated"):
        self.templates_dir = Path(templates_dir)
        self.output_dir = Path(output_dir)
        
        # Try to create directory and verify writability
        try:
            self.output_dir.mkdir(parents=True, exist_ok=True)
            # Test write permission
            test_path = self.output_dir / ".write_test"
            with open(test_path, "w") as f:
                f.write("test")
            test_path.unlink()
        except (OSError, PermissionError):
            # Fallback to /tmp if read-only or permission denied
            self.output_dir = Path("/tmp/generated_docs")
            self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_offer_letter(self, consultant_data):
        """Generate personalized offer letter from Solutions template"""
        
        # Load template
        template_path = self.templates_dir / "offer_letter_template.docx"
        doc = Document(template_path)
        
        # Calculate dates
        today = datetime.now().strftime("%B %d, %Y")
        response_deadline = (datetime.now() + timedelta(days=3)).strftime("%B %d, %Y")
        
        # Prepare replacement data
        replacements = {
            '{{date}}': today,
            '{{consultant_name}}': consultant_data['name'],
            '{{consultant_email}}': consultant_data.get('email', ''),
            '{{consultant_first_name}}': consultant_data['name'].split()[0],
            '{{position_title}}': consultant_data['position'],
            '{{employment_type_full}}': 'X' if consultant_data.get('employment_type') == 'Full-time' else '☐',
            '{{employment_type_part}}': '☐' if consultant_data.get('employment_type') == 'Full-time' else 'X',
            '{{exempt_yes}}': 'X' if consultant_data.get('exempt_status') == 'Exempt' else '☐',
            '{{exempt_no}}': '☐' if consultant_data.get('exempt_status') == 'Exempt' else 'X',
            '{{location_onsite}}': '☐',
            '{{location_remote}}': 'X' if 'Remote' in consultant_data.get('work_location', '') else '☐',
            '{{location_hybrid}}': 'X' if 'Hybrid' in consultant_data.get('work_location', '') else '☐',
            '{{location_detail}}': consultant_data.get('location_detail', 'Portland, Maine'),
            '{{manager_name}}': consultant_data.get('manager', 'Debbie Murray'),
            '{{start_date}}': consultant_data['start_date'],
            '{{end_date}}': consultant_data.get('end_date', 'N/A'),
            '{{pay_rate}}': consultant_data.get('pay_rate', ''),
            '{{pay_hourly}}': 'X' if consultant_data.get('pay_frequency') == 'Hourly' else '☐',
            '{{pay_yearly}}': '☐' if consultant_data.get('pay_frequency') == 'Hourly' else 'X',
            '{{pay_schedule}}': consultant_data.get('pay_schedule', 'X Monthly'),
            '{{min_hours}}': consultant_data.get('min_hours', '32'),
            '{{max_hours}}': consultant_data.get('max_hours', '40'),
            '{{response_deadline}}': response_deadline,
            '{{hiring_manager_name}}': consultant_data.get('hiring_manager', 'Debbie Murray'),
            '{{hiring_manager_title}}': consultant_data.get('hiring_manager_title', 'President'),
            '{{company_name}}': 'Solutions Project Management, LLC'
        }
        
        # Replace text in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)
        
        # Replace text in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replacements.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)
        
        # Save document
        output_filename = f"Offer_Letter_{consultant_data['name'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        output_path = self.output_dir / output_filename
        doc.save(output_path)
        
        return str(output_path)
    
    def generate_job_description(self, consultant_data, responsibilities, requirements):
        """Generate job description document"""
        
        doc = Document()
        
        # Title
        doc.add_heading(f'Job Description: {consultant_data["position"]}', 0)
        
        # Position Overview
        doc.add_heading('Position Overview', level=1)
        doc.add_paragraph(f'Position: {consultant_data["position"]}')
        doc.add_paragraph(f'Department: {consultant_data.get("department", "Consulting")}')
        doc.add_paragraph(f'Reports To: {consultant_data["manager"]}')
        doc.add_paragraph(f'Employment Type: {consultant_data["employment_type"]}')
        
        # Responsibilities
        doc.add_heading('Key Responsibilities', level=1)
        for resp in responsibilities:
            doc.add_paragraph(resp, style='List Bullet')
        
        # Requirements
        doc.add_heading('Requirements', level=1)
        for req in requirements:
            doc.add_paragraph(req, style='List Bullet')
        
        # Save
        output_filename = f"Job_Description_{consultant_data['position'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        output_path = self.output_dir / output_filename
        doc.save(output_path)
        
        return str(output_path)
    
    def generate_onboarding_checklist(self, consultant_data):
        """Generate onboarding checklist"""
        
        doc = Document()
        
        # Title
        doc.add_heading(f'Onboarding Checklist: {consultant_data["name"]}', 0)
        doc.add_paragraph(f'Start Date: {consultant_data["start_date"]}')
        doc.add_paragraph('')
        
        # Checklist items
        checklist_items = [
            ('Before Start Date', [
                'Review and sign offer letter',
                'Complete background check authorization',
                'Submit signed offer letter'
            ]),
            ('First Day', [
                'Complete I-9 (bring required documents)',
                'Complete W-4 tax withholding form',
                'Set up direct deposit',
                'Review employee handbook',
                'Complete emergency contact information',
                'Sign confidentiality agreement',
                'Receive company equipment (laptop, phone, etc.)'
            ]),
            ('First Week', [
                'Complete benefits enrollment',
                'Set up email and system access',
                'Schedule orientation meetings',
                'Review company policies',
                'Meet team members'
            ]),
            ('First Month', [
                'Complete required training modules',
                'Set initial performance goals',
                'Schedule 30-day check-in with manager'
            ])
        ]
        
        for section_title, items in checklist_items:
            doc.add_heading(section_title, level=1)
            for item in items:
                doc.add_paragraph(f'☐ {item}', style='List Bullet')
            doc.add_paragraph('')
        
        # Save
        output_filename = f"Onboarding_Checklist_{consultant_data['name'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        output_path = self.output_dir / output_filename
        doc.save(output_path)
        
        return str(output_path)